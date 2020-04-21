"""
使用docx 处理word文档
@Time : 2020/4/7 16:59
@Author: zhangqian
"""
from docx import Document

doc = Document('./res/1.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].runs[-1].text)
print(doc.paragraphs[1].runs[0].text)